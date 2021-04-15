import openpyxl
import docx
from docx.shared import Inches
import os

#Interestingly, the file path will immediately appear underneath the Finder window.
# Now, Control + Click the file you want the location for
# Hold the Option key
# You will see the new command that has appeared in the context menu — Copy …. as Pathname

path='/Users/chorongbali/Desktop/PythonDocs/PythonExcel.xlsx'
wb = openpyxl.load_workbook(path)
sheetnames = wb.get_sheet_names()
print(sheetnames)
# wb.create_sheet("2nd Sheet", 1)
# wb.save(path)

print(f"The file located at {wb}")

# worksheets represents the tab of sheets in that excel file
ws = wb.worksheets[0]
# col1 = ws["A1"].value
# col2 = (ws["A2"].value)
# print(ws["A2"].value)
# for col in range(1, 4):
#     print(ws.cell(column=col, row=2).value)
#     print(ws.cell(column=col, row=2).value)
#     print(ws.cell(column=col, row=3).value)
#
# if col1 == 'NAME':
#     print(col2)

# for col in range(1, 5):
#     if col != '':
#         print(ws.cell(column=col, row=2).value)
        # print(ws.cell(column=col, row=3).value)

# for x in range (1,5,1):
#     print(x,ws.cell(row=x,column=1).value,
#     x,ws.cell(row=x,column=2).value,
#     x,ws.cell(row=x,column=3).value,
#     x,ws.cell(row=x,column=4).value)
    # print(x,ws.cell(row=x,column=3).value,
    # print(x,ws.cell(row=x,column=4).value))

listOfNames = []

for x in range(2,10):
    listOfNames.append((ws.cell(row=x, column=1).value))
print(listOfNames)


# DOCX
pathdocx = '/Users/chorongbali/Desktop/PythonDocs/PythonDocx.docx'
doc = docx.Document(pathdocx)
all_para = doc.paragraphs
print(all_para)
print(len(doc.paragraphs))

for para in all_para:
    print(para.text)

for paragraph in doc.paragraphs:
    if 'Bali' in paragraph.text:
        print(f"old text is {paragraph.text}")
        paragraph.text = 'new bali? hehe'
        print(f"new text is {paragraph.text}")
# bali = (listOfNames[1])
# doc.add_paragraph(bali)
# doc.add_paragraph("This is the 2nd paragraph!")


# HANDLING TABLES WITH DOCX

print(doc.tables)
# table = doc.add_table(rows=2, cols=2, style='Table Grid')

table1 = doc.tables[0]
table2 = doc.tables[1]
table3 = doc.tables[2]
print(table1)
print(listOfNames)

table = (table1)
heading_cells = table.rows[2].cells
heading_cells[2].text = (listOfNames[0])
print(heading_cells)

# for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print(cell.text)
            # for paragraph in cell.paragraphs:
                # print(paragraph.text)
                # if 'MAPPING TABLE' in paragraph.text:
                #     print("YOU FOUND THE MAPPING TABLE")



for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if 'sea' in paragraph.text:
                    paragraph.text = paragraph.text.replace("sea", "ocean")

for table in doc.tables:
    for row in table.rows:
        if 'MAPPING TABLE' in table.rows:
            print("NICE U FOUND IT")
        # if table.rows[1] == 'MAPPING TABLE':
        #     print("U FOIND IT BRO")

if table.rows[1] == 'MAPPING TABLE':
    print("YOU FOUND THE MAPPING TABLE!")

doc.save(pathdocx)





